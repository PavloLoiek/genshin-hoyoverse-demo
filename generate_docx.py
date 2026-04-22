#!/usr/bin/env python3
"""
Generate DOCX files from content JSON files for ACME Inn demo.
Reads from documents_content/ and outputs to documents/ and manual_upload/.
"""

import json
import os
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = Path(__file__).parent
CONTENT_DIR = BASE_DIR / "documents_content"
DOCS_DIR = BASE_DIR / "documents"
MANUAL_UPLOAD_DIR = BASE_DIR / "manual_upload"

# Brand colors
ACME_BLUE = RGBColor(0x1a, 0x3c, 0x6e)
ACME_DARK = RGBColor(0x1a, 0x23, 0x32)
ACME_MUTED = RGBColor(0x5a, 0x6a, 0x7a)


def setup_styles(doc):
    """Configure document styles for professional appearance."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = ACME_DARK

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = ACME_BLUE
        if level == 1:
            heading_style.font.size = Pt(20)
            heading_style.font.bold = True
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        else:
            heading_style.font.size = Pt(13)
            heading_style.font.bold = True


def add_title_page(doc, title, customer_name, doc_type):
    """Add a professional title page."""
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACME_BLUE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{'_' * 60}")
    run.font.color.rgb = ACME_MUTED
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(customer_name)
    run.font.size = Pt(16)
    run.font.color.rgb = ACME_DARK
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_labels = {
        'faq': 'Frequently Asked Questions',
        'sop': 'Standard Operating Procedure',
        'policy': 'Business Policy Document',
        'training': 'Training Guide',
        'memo': 'Internal Memorandum',
        'reference': 'Reference Document',
        'misc': 'Reference Document',
    }
    run = p.add_run(type_labels.get(doc_type, 'Document'))
    run.font.size = Pt(12)
    run.font.color.rgb = ACME_MUTED
    run.font.italic = True

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — For Internal Use Only")
    run.font.size = Pt(9)
    run.font.color.rgb = ACME_MUTED
    run.font.italic = True

    doc.add_page_break()


def generate_faq(doc, data):
    """Generate FAQ document with Q&A pairs."""
    for section in data.get('sections', []):
        heading = section.get('heading', '').replace('_', ' ')
        doc.add_heading(heading, level=1)

        for qa in section.get('qa_pairs', []):
            p = doc.add_paragraph()
            run = p.add_run(f"Q: {qa['q']}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = ACME_BLUE

            p = doc.add_paragraph()
            run = p.add_run(f"A: {qa['a']}")
            run.font.size = Pt(11)

            doc.add_paragraph()


def generate_sop(doc, data):
    """Generate SOP document with numbered procedures."""
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        doc.add_heading(heading, level=1)

        content = section.get('content', '')
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            if para_text and para_text[0].isdigit() and '.' in para_text[:4]:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                p.paragraph_format.left_indent = Cm(1)
            elif para_text.startswith('- '):
                p = doc.add_paragraph(para_text[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)


def generate_policy(doc, data):
    """Generate policy document with formal sections."""
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        if heading:
            doc.add_heading(heading, level=1)

        content = section.get('content', '')
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            if para_text and para_text[0].isdigit() and '.' in para_text[:5]:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                p.paragraph_format.left_indent = Cm(0.5)
            elif para_text.startswith('- '):
                p = doc.add_paragraph(para_text[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)


def generate_training(doc, data):
    """Generate training guide document."""
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        doc.add_heading(heading, level=1)

        content = section.get('content', '')
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            if para_text and para_text[0].isdigit() and '.' in para_text[:4]:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                p.paragraph_format.left_indent = Cm(1)
            elif para_text.startswith('- '):
                p = doc.add_paragraph(para_text[2:], style='List Bullet')
            elif para_text.startswith('Tip:') or para_text.startswith('Note:'):
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = ACME_BLUE
            else:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)


def generate_memo(doc, data):
    """Generate memo/misc document."""
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        if heading:
            doc.add_heading(heading, level=1)

        content = section.get('content', '')
        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue

            if ':' in para_text and para_text.index(':') < 30 and not para_text[0].isdigit():
                parts = para_text.split(':', 1)
                p = doc.add_paragraph()
                run = p.add_run(parts[0] + ':')
                run.bold = True
                run.font.size = Pt(11)
                run2 = p.add_run(parts[1])
                run2.font.size = Pt(11)
            elif para_text.startswith('- '):
                p = doc.add_paragraph(para_text[2:], style='List Bullet')
            elif ' -> ' in para_text or ' — ' in para_text:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                run.font.name = 'Consolas'
            else:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(11)


def process_json_file(json_path, output_dir, is_misc=False):
    """Process a single JSON content file and generate DOCX."""
    with open(json_path) as f:
        data = json.load(f)

    doc_type = data.get('type', 'misc')
    title = data.get('title', json_path.stem)
    customer_name = data.get('customer_name', 'ACME Inn')

    doc = Document()
    setup_styles(doc)
    add_title_page(doc, title, customer_name, doc_type)

    generators = {
        'faq': generate_faq,
        'sop': generate_sop,
        'policy': generate_policy,
        'training': generate_training,
        'memo': generate_memo,
        'reference': generate_memo,
    }

    generator = generators.get(doc_type, generate_memo)
    generator(doc, data)

    output_name = json_path.stem + '.docx'
    if is_misc:
        output_path = output_dir / 'misc' / output_name
    else:
        output_path = output_dir / output_name

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  Generated: {output_path.relative_to(BASE_DIR)}")
    return output_path


def copy_to_manual_upload(docx_path, json_path, is_misc=False):
    """Copy DOCX to the appropriate manual_upload subfolder."""
    stem = json_path.stem
    doc_type_from_name = ''
    if stem.endswith('_sop') or stem.endswith('_policy'):
        doc_type_from_name = 'policy_sop'
    elif stem == 'faq_document':
        doc_type_from_name = 'faq'
    elif stem == 'general_training_guide':
        doc_type_from_name = 'policy_sop'
    elif is_misc:
        doc_type_from_name = 'misc'
    else:
        doc_type_from_name = 'policy_sop'

    dest_dirs = {
        'faq': MANUAL_UPLOAD_DIR / '1_faq',
        'policy_sop': MANUAL_UPLOAD_DIR / '3_policies_and_sops',
        'misc': MANUAL_UPLOAD_DIR / '5_misc_docs',
    }

    dest_dir = dest_dirs.get(doc_type_from_name)
    if dest_dir:
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = dest_dir / docx_path.name
        shutil.copy2(str(docx_path), str(dest_path))
        print(f"  Copied to: {dest_path.relative_to(BASE_DIR)}")


def main():
    print("=" * 60)
    print("ACME Inn — DOCX Document Generator")
    print("=" * 60)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    (DOCS_DIR / 'misc').mkdir(parents=True, exist_ok=True)

    generated = []

    print("\nProcessing main documents:")
    for json_file in sorted(CONTENT_DIR.glob('*.json')):
        docx_path = process_json_file(json_file, DOCS_DIR, is_misc=False)
        copy_to_manual_upload(docx_path, json_file, is_misc=False)
        generated.append(docx_path)

    misc_dir = CONTENT_DIR / 'misc'
    if misc_dir.exists():
        print("\nProcessing misc documents:")
        for json_file in sorted(misc_dir.glob('*.json')):
            docx_path = process_json_file(json_file, DOCS_DIR, is_misc=True)
            copy_to_manual_upload(docx_path, json_file, is_misc=True)
            generated.append(docx_path)

    print(f"\nTotal DOCX files generated: {len(generated)}")
    print("=" * 60)


if __name__ == '__main__':
    main()
